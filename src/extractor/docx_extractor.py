from docx import Document
from pathlib import Path

from src.utils.logging_utils import LoggingService

logger = LoggingService.get_logger("extractor.docx")

class DocxExtractor:
    """Classe responsável por extrair texto de arquivos DOCX."""

    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """Extrai texto de um arquivo DOCX usando a biblioteca python-docx."""
        try:
            document = Document(str(file_path))

            logger.info(f"Começando extração de documento DOCX: {file_path}")

            document_rows = []
            paragraph_count = 0
            table_rows_count = 0

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    document_rows.append(text)
                    paragraph_count += 1

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        document_rows.append(" | ".join(cells))
                        table_rows_count += 1

            full_text = "\n".join(document_rows)

            logger.info(
                f"Extração finalizada: {file_path}\n"
                f"paragrafos={paragraph_count}, linhas_tabelas={table_rows_count}, "
                f"total_linhas={len(document_rows)}"
            )

            if not full_text.strip():
                logger.warning(f"Nenhum texto extraído de: {file_path}")

            return full_text
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX {file_path}: {e}")
            return ""